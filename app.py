import streamlit as st
import pandas as pd
import plotly.express as px
import requests
import math
import io
from datetime import datetime, date, timedelta

# ── CONFIGURATION ─────────────────────────────────────────────────────────────

# ⚠️ ตรวจสอบ SCRIPT_URL ให้ถูกต้อง
SCRIPT_URL = "https://script.google.com/macros/s/AKfycbzKjgLJ7yRHLkDpCZejbmWEDBQcvyd-YZmeS7WMMYBVkKkkyhckElmRVoE1NpHNenX7NA/exec"

MONTHLY_FTE_MINUTES = 8_750
ANNUAL_FTE_MINUTES  = 105_000

# รายการงาน
TASKS = {
    1:  "ควบคุมคุณภาพและจัดทำรายงานการควบคุมคุณภาพเครื่องมือทางรังสีวินิจฉัย รวมถึงติดตาม ควบคุม ประจำวัน สัปดาห์ เดือน ครึ่งปี ปี",
    2:  "ติดตาม ตรวจสอบ ความผิดปกติของเครื่องมือทางรังสีวินิจฉัยร่วมกับนักรังสีการแพทย์ และ วิศวกร",
    3:  "วัดและประเมินปริมาณรังสีกระเจิง และ ปริมาณรังสีรั่วไหล รวมถึงจัดทำรายงาน",
    4:  "จัดทำค่าปริมาณรังสีอ้างอิง (DRL) และจัดทำโปรโตคอลในการกำหนดปริมาณรังสีที่ผู้ป่วยได้รับ",
    5:  "ดำเนินการเกี่ยวกับการวัดปริมาณรังสีส่วนบุคคล (OSL/TLD) บันทึกรายงาน และหาทางรับมือหากรังสีเกินมาตรฐาน",
    6:  "ตรวจสอบคุณภาพอุปกรณ์กำบังรังสี เช่น เสื้อตะกั่ว thyroid shield",
    7:  "คำนวณปริมาณรังสีในผู้ป่วยตั้งครรภ์ รวมถึงผู้ป่วยที่ได้รับปริมาณรังสีสูงเกินขีดจำกัด",
    8:  "ตรวจสอบแรกรับ และประเมินประสิทธิภาพการใช้งานของเครื่องมือทางรังสีหลังการติดตั้ง ซ่อมแซม (Acceptance Test)",
    9:  "จัดทำและเข้าร่วมในการกำหนดรายละเอียดคุณลักษณะเฉพาะ (TOR/Spec) เพื่อการจัดซื้อเครื่องมือ",
    10: "จัดทำแผนการป้องกันอันตรายทางรังสี แผนฉุกเฉิน และแผนอุบัติเหตุทางรังสี",
    11: "ให้คำแนะนำการป้องกันอันตรายจากรังสี รวมถึงแผนฉุกเฉินทางรังสี",
    12: "ศึกษา ค้นคว้า วิเคราะห์ วิจัยทางฟิสิกส์การแพทย์",
    13: "ดำเนินการสอบเทียบมาตรฐานอุปกรณ์ที่ใช้ในการวัดปริมาณรังสี (Calibration)",
    14: "จัดทำรายงานการขออนุญาตการมีไว้ในครอบครองและการใช้งานเครื่องกำเนิดรังสี",
    15: "ตรวจสอบการได้รับรังสีกรณีเกิดเหตุใดๆ กับผู้ป่วยหลังการตรวจ",
    16: "ประชุมหัวข้ออื่นๆที่เกี่ยวข้องกับทางรังสีวินิจฉัย"
}

TASK_LABELS = [f"{k}. {v[:60]}…" if len(v) > 60 else f"{k}. {v}" for k, v in TASKS.items()]
TASK_NO_FROM_LABEL = {label: no for no, label in zip(TASKS.keys(), TASK_LABELS)}

THAI_MONTHS = {
    1: "มกราคม", 2: "กุมภาพันธ์", 3: "มีนาคม", 4: "เมษายน",
    5: "พฤษภาคม", 6: "มิถุนายน", 7: "กรกฎาคม", 8: "สิงหาคม",
    9: "กันยายน", 10: "ตุลาคม", 11: "พฤศจิกายน", 12: "ธันวาคม",
}

# ── PAGE CONFIG ───────────────────────────────────────────────────────────────

st.set_page_config(page_title="Medical Physicist Workload", page_icon="⚛️", layout="wide")

st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Sarabun:wght@300;400;600;700&display=swap');
    html, body, [class*="css"] { font-family: 'Sarabun', sans-serif; }
    [data-testid="stSidebar"] { background: linear-gradient(180deg, #0f2027 0%, #203a43 50%, #2c5364 100%); }
    .section-header { 
        background: #009ef7; 
        color: white; 
        padding: 12px 20px; 
        border-radius: 10px; 
        margin: 10px 0; 
        font-weight: 700;
        display: flex;
        align-items: center;
        gap: 10px;
    }
    .fte-box { background: linear-gradient(135deg, #0f2027, #2c5364); color: #bae6fd; padding: 20px; border-radius: 12px; text-align: center; margin: 15px 0; }
    .fte-box span.big { font-size: 2.2rem; color: #38bdf8; font-weight: 700; }
    .fte-box span.med { font-size: 1.6rem; color: #fca5a5; font-weight: 700; }
    footer { visibility: hidden; }
</style>
""", unsafe_allow_html=True)

# ── HELPER FUNCTIONS ──────────────────────────────────────────────────────────

def format_gas_time(time_val):
    if pd.isna(time_val) or not time_val or str(time_val).strip() == "-" or str(time_val).lower() == 'nan': 
        return "-"
    time_str = str(time_val)
    if "T" in time_str or "Z" in time_str:
        try:
            dt = pd.to_datetime(time_str)
            if dt.tz is not None:
                dt = dt.tz_convert('Asia/Bangkok')
            else:
                dt = dt + pd.Timedelta(hours=7)
            return dt.strftime("%H:%M")
        except: pass
    return time_str[:5]

@st.cache_data(ttl=60)
def fetch_all_data():
    try:
        r = requests.get(SCRIPT_URL, timeout=15)
        return r.json() if r.status_code == 200 else {"logs": [], "staff": []}
    except:
        return {"logs": [], "staff": []}

def invalidate_cache():
    fetch_all_data.clear()

def generate_word_report(df_report, report_type, period_text):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = Document()
    title_heading = doc.add_heading('รายงานสรุปภาระงาน (Workload Report)', 0)
    title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"หน่วยงาน: สาขารังสีวินิจฉัย โรงพยาบาลสงขลานครินทร์")
    doc.add_paragraph(f"ประเภทรายงาน: {report_type}")
    doc.add_paragraph(f"ช่วงเวลา: {period_text}")
    
    thai_now = datetime.now() + timedelta(hours=7)
    doc.add_paragraph(f"วันที่พิมพ์รายงาน: {thai_now.strftime('%d/%m/%Y %H:%M')}")
    
    total_mins = df_report["Minutes"].sum()
    fte = round(total_mins / (ANNUAL_FTE_MINUTES if "ปี" in report_type else MONTHLY_FTE_MINUTES), 2)
    req_staff = math.ceil(fte) if fte > 0 else 0
    
    doc.add_paragraph(f"สรุปเวลารวม: ปฏิบัติงานทั้งสิ้น {total_mins:,.0f} นาที (คิดเป็นภาระงาน {fte} FTE)")
    staff_paragraph = doc.add_paragraph(f"📌 อัตรากำลังนักฟิสิกส์การแพทย์ที่เหมาะสมกับภาระงาน: {req_staff} คน")
    staff_paragraph.runs[0].bold = True
    
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    headers = ['วันที่', 'เวลา', 'ผู้ปฏิบัติงาน', 'ชื่องาน', 'รายละเอียด', 'นาที']
    for i, txt in enumerate(headers):
        hdr[i].text = txt

    for _, row in df_report.sort_values(by=["Date", "Start_Time"]).iterrows():
        cells = table.add_row().cells
        cells[0].text = pd.to_datetime(row['Date']).strftime('%d/%m/%Y')
        s_time = format_gas_time(row.get('Start_Time', '-'))
        e_time = format_gas_time(row.get('End_Time', '-'))
        cells[1].text = f"{s_time}-{e_time}"
        cells[2].text = str(row['Name'])
        cells[3].text = str(row['Task_Name'])
        
        det = str(row.get('Details', ''))
        cells[4].text = det if det != 'nan' and det.strip() != '' else '-'
        cells[5].text = str(row['Minutes'])
        
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# ── APP LOGIC ─────────────────────────────────────────────────────────────────

with st.sidebar:
    st.markdown("## ⚛️ Medical Physicist\n### Workload Manager")
    st.markdown("---")
    tab_choice = st.radio("เมนูหลัก", ["📝 บันทึกภาระงาน", "📅 ประวัติงาน & ปฏิทิน", "📊 Dashboard", "⚙️ ออกรายงาน & จัดการข้อมูล"])
    st.markdown("---")
    st.caption("v1.15 — Daily Report Export Added")

data = fetch_all_data()
staff_list = data.get("staff", [])
df = pd.DataFrame(data.get("logs", []))

if not df.empty:
    df["Minutes"] = pd.to_numeric(df["Minutes"], errors="coerce").fillna(0)
    parsed_dates = pd.to_datetime(df["Date"], errors='coerce')
    if parsed_dates.dt.tz is not None:
        df["Date"] = parsed_dates.dt.tz_convert('Asia/Bangkok').dt.date
    else:
        df["Date"] = (parsed_dates + pd.Timedelta(hours=7)).dt.date

# --- TAB 1: ENTRY ---
if tab_choice == "📝 บันทึกภาระงาน":
    st.markdown('<div class="section-header">📝 บันทึกภาระงานประจำวัน</div>', unsafe_allow_html=True)
    if not staff_list: st.warning("กรุณาเพิ่มชื่อเจ้าหน้าที่ในเมนู 'จัดการข้อมูล' ก่อน"); st.stop()

    with st.form("entry_form", clear_on_submit=True):
        c1, c2 = st.columns(2)
        with c1:
            e_date = st.date_input("วันที่", value=date.today())
            e_name = st.selectbox("ชื่อผู้ปฏิบัติงาน", options=staff_list)
            e_desc = st.text_input("รายละเอียดเพิ่มเติม (Details)")
        with c2:
            e_task_label = st.selectbox("รายการงาน", options=TASK_LABELS)
            tc1, tc2 = st.columns(2)
            with tc1: s_time = st.time_input("เวลาเริ่ม", value=datetime.strptime("08:30", "%H:%M").time())
            with tc2: e_time = st.time_input("เวลาสิ้นสุด", value=datetime.strptime("09:00", "%H:%M").time())
        
        if st.form_submit_button("💾 บันทึกข้อมูล", use_container_width=True, type="primary"):
            dt_s, dt_e = datetime.combine(e_date, s_time), datetime.combine(e_date, e_time)
            if dt_e < dt_s: dt_e += timedelta(days=1)
            mins = int((dt_e - dt_s).total_seconds() / 60)
            
            task_no = TASK_NO_FROM_LABEL[e_task_label]
            payload = {
                "action": "append_log",
                "data": [datetime.now().strftime("%Y-%m-%d %H:%M:%S"), str(e_date), e_name, 
                         task_no, TASKS[task_no], e_desc, mins, s_time.strftime("%H:%M"), e_time.strftime("%H:%M")]
            }
            if requests.post(SCRIPT_URL, json=payload).status_code in [200, 302]:
                st.success(f"บันทึกสำเร็จ! ({mins} นาที)"); invalidate_cache(); st.rerun()

# --- TAB 2: HISTORY & CALENDAR ---
elif tab_choice == "📅 ประวัติงาน & ปฏิทิน":
    st.markdown('<div class="section-header">📅 รายการภาระงานย้อนหลัง</div>', unsafe_allow_html=True)
    if df.empty: st.info("ไม่มีข้อมูล"); st.stop()
    
    temp_df = df.copy()
    temp_df["Time"] = temp_df.apply(lambda x: f"{format_gas_time(x.get('Start_Time','-'))} - {format_gas_time(x.get('End_Time','-'))}", axis=1)
    
    view_df = temp_df[["Date", "Time", "Name", "Task_Name", "Details", "Minutes"]].sort_values(["Date", "Time"], ascending=[False, False])
    view_df["Details"] = view_df["Details"].fillna("-").replace("nan", "-").replace("", "-")
    
    st.dataframe(view_df, use_container_width=True, hide_index=True)

    st.markdown("---")
    st.subheader("📆 ปฏิทินภาระงาน")
    try:
        from streamlit_calendar import calendar
        events = [{"title": f"{r['Name']}: {r['Minutes']}m", "start": str(r["Date"]), "end": str(r["Date"]), "color": "#009ef7"} for _, r in df.iterrows()]
        calendar(events=events, options={"initialView": "dayGridMonth", "locale": "th"}, key="cal")
    except: st.caption("ติดตั้ง streamlit-calendar เพื่อดูปฏิทิน")

    st.markdown("---")
    st.subheader("🗑️ ลบรายการ")
    df_with_idx = df.copy()
    df_with_idx["idx"] = df_with_idx.index + 2
    row_map = {f"แถว {r['idx']} | {r['Date']} | {r['Name']}": r['idx'] for _, r in df_with_idx.iterrows()}
    target = st.selectbox("เลือกรายการที่จะลบ", options=list(row_map.keys()))
    if st.button("❌ ยืนยันการลบ"):
        if requests.post(SCRIPT_URL, json={"action": "delete_log", "row_index": row_map[target]}).status_code in [200, 302]:
            st.success("ลบแล้ว"); invalidate_cache(); st.rerun()

# --- TAB 3: DASHBOARD ---
elif tab_choice == "📊 Dashboard":
    st.markdown('<div class="section-header">📊 Dashboard & FTE Analysis</div>', unsafe_allow_html=True)
    if df.empty: st.info("ยังไม่มีข้อมูล"); st.stop()
    
    view_type = st.radio("เลือกมุมมองการประเมินภาระงาน", ["📅 สรุปรายเดือน", "📆 สรุปรายปี"], horizontal=True)
    st.markdown("<br>", unsafe_allow_html=True)

    yrs = sorted(pd.to_datetime(df["Date"]).dt.year.unique(), reverse=True)
    
    if view_type == "📅 สรุปรายเดือน":
        c1, c2 = st.columns(2)
        with c1: sel_y = st.selectbox("เลือกปี", yrs)
        with c2: sel_m = st.selectbox("เลือกเดือน", range(1, 13), format_func=lambda x: THAI_MONTHS[x])
        
        target_df = df[(pd.to_datetime(df["Date"]).dt.year == sel_y) & (pd.to_datetime(df["Date"]).dt.month == sel_m)]
        total = target_df["Minutes"].sum()
        fte = round(total / MONTHLY_FTE_MINUTES, 2)
        req_staff = math.ceil(fte) if fte > 0 else 0
        title_text = f"สรุปภาระงานเดือน {THAI_MONTHS[sel_m]} {sel_y}"
        
    else:
        sel_y = st.selectbox("เลือกปี", yrs)
        target_df = df[pd.to_datetime(df["Date"]).dt.year == sel_y]
        total = target_df["Minutes"].sum()
        fte = round(total / ANNUAL_FTE_MINUTES, 2)
        req_staff = math.ceil(fte) if fte > 0 else 0
        title_text = f"สรุปภาระงานประจำปี {sel_y}"
    
    st.markdown(f'''
    <div class="fte-box">
        {title_text}<br>
        <span class="big">{fte}</span> FTE <span style="font-size: 1rem;">({total:,.0f} นาที)</span><br>
        <div style="margin-top: 10px; padding-top: 10px; border-top: 1px solid rgba(255,255,255,0.1);">
            📌 อัตรากำลังนักฟิสิกส์การแพทย์ที่ต้องการ: <span class="med">{req_staff}</span> คน
        </div>
    </div>
    ''', unsafe_allow_html=True)
    
    if not target_df.empty:
        st.plotly_chart(px.bar(target_df.groupby("Name")["Minutes"].sum().reset_index(), x="Name", y="Minutes", color="Name", title="สัดส่วนภาระงานรายบุคคล (นาที)"), use_container_width=True)

# --- TAB 4: EXPORT & SETTINGS ---
elif tab_choice == "⚙️ ออกรายงาน & จัดการข้อมูล":
    st.markdown('<div class="section-header">⚙️ จัดการข้อมูล & ออกรายงาน Word</div>', unsafe_allow_html=True)
    
    st.subheader("👤 จัดการรายชื่อเจ้าหน้าที่")
    c_s1, c_s2 = st.columns(2)
    with c_s1:
        new_n = st.text_input("เพิ่มชื่อใหม่")
        if st.button("💾 เพิ่มชื่อ") and new_n:
            requests.post(SCRIPT_URL, json={"action": "add_staff", "name": new_n})
            invalidate_cache(); st.rerun()
    with c_s2:
        if staff_list:
            del_n = st.selectbox("เลือกชื่อที่จะลบ", staff_list)
            if st.button("🗑️ ลบรายชื่อ"):
                requests.post(SCRIPT_URL, json={"action": "delete_staff", "name": del_n})
                invalidate_cache(); st.rerun()

    st.markdown("---")
    
    st.subheader("📥 ออกรายงาน Workload (Word .docx)")
    # เพิ่มตัวเลือก "รายวัน" ใน Radio button
    etype = st.radio("เลือกช่วงเวลา", ["รายวัน", "สัปดาห์/กำหนดเอง", "รายเดือน", "รายปี"], horizontal=True)
    
    if not df.empty:
        d_exp = pd.DataFrame()
        title, p_str = "", ""
        
        if etype == "รายวัน":
            target_day = st.date_input("เลือกวันที่", value=date.today())
            d_exp = df[df['Date'] == target_day]
            title, p_str = "รายงานประจำวัน", f"วันที่ {target_day.strftime('%d/%m/%Y')}"
            
        elif "สัปดาห์" in etype:
            sd = st.date_input("เริ่ม", value=date.today()-timedelta(7))
            ed = st.date_input("จบ", value=date.today())
            d_exp = df[(df['Date'] >= sd) & (df['Date'] <= ed)]
            title, p_str = "รายงานประจำสัปดาห์", f"{sd.strftime('%d/%m/%Y')} ถึง {ed.strftime('%d/%m/%Y')}"
            
        elif "เดือน" in etype:
            yrs_avail = sorted(pd.to_datetime(df["Date"]).dt.year.unique(), reverse=True)
            ey = st.selectbox("ปี", yrs_avail, key="ey")
            em = st.selectbox("เดือน", range(1,13), format_func=lambda x: THAI_MONTHS[x], key="em")
            d_exp = df[(pd.to_datetime(df["Date"]).dt.year == ey) & (pd.to_datetime(df["Date"]).dt.month == em)]
            title, p_str = "รายงานรายเดือน", f"{THAI_MONTHS[em]} {ey}"
            
        else:
            yrs_avail = sorted(pd.to_datetime(df["Date"]).dt.year.unique(), reverse=True)
            ey2 = st.selectbox("ปี", yrs_avail, key="ey2")
            d_exp = df[pd.to_datetime(df["Date"]).dt.year == ey2]
            title, p_str = "รายงานรายปี", f"พ.ศ. {ey2 + 543}"

        if d_exp.empty: 
            st.info("⚠️ ไม่พบข้อมูลในช่วงเวลาที่เลือก")
        else:
            f_bytes = generate_word_report(d_exp, title, p_str)
            st.download_button(
                label=f"📥 ดาวน์โหลดไฟล์ Word ({title})", 
                data=f_bytes, 
                file_name=f"Workload_{etype}_{datetime.now().strftime('%Y%m%d')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
